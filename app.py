"""
Legal Document Analyzer App
"""

import streamlit as st
import os
import logging
import json
import re
import tempfile
from typing import List, Dict, Any, Tuple, Optional
import numpy as np
import faiss
import torch
from sentence_transformers import SentenceTransformer

# Import local modules
from utils import save_uploaded_file, format_bytes, get_temp_file_path, clean_filename, generate_document_id
from pdf_processor import extract_text_from_pdf, clean_text, split_into_chunks
from metadata_extractor import MetadataExtractor
from embedding_service import create_embeddings, create_index, search_similar
from sambanova_api import get_llama_response
from summary_generator import generate_document_summary, compare_documents

# Setup page config right at the start - before any other st commands
st.set_page_config(page_title="Document Chat", layout="wide")

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    datefmt='%Y-%m-%d %H:%M:%S'
)
logger = logging.getLogger(__name__)

# Initialize metadata extractor
metadata_extractor = MetadataExtractor()

# Custom CSS to improve app appearance
custom_css = """
<style>
.document-card {
    background-color: #333;
    padding: 1rem;
    border-radius: 5px;
    margin-bottom: 1rem;
    border-left: 3px solid #4e8cff;
}
.metadata-label {
    font-weight: bold;
    color: #aaa;
}
.stButton>button {
    width: 100%;
}
.chat-message-user {
    background-color: #2c5282; /* Darker blue for user messages */
    padding: 0.75rem;
    border-radius: 15px 15px 0 15px;
    margin-bottom: 0.5rem;
    color: white;
}
.chat-message-ai {
    background-color: #333; /* Darker gray for AI messages */
    padding: 0.75rem;
    border-radius: 15px 15px 15px 0;
    margin-bottom: 0.5rem;
    color: #eee;
}
/* Updated tab styling for dark theme */
div.stTabs {
    background-color: #1e1e1e;
}
div.stTabs [data-baseweb="tab-list"] {
    gap: 2px;
    background-color: #1e1e1e;
}
div.stTabs [data-baseweb="tab"] {
    height: 50px;
    white-space: pre-wrap;
    background-color: #2d2d2d;
    border-radius: 4px 4px 0 0;
    gap: 1px;
    padding-top: 10px;
    padding-bottom: 10px;
    color: #aaa;
}
div.stTabs [aria-selected="true"] {
    background-color: #4e8cff;
    color: white;
}
/* Improve tab hover effect */
div.stTabs [data-baseweb="tab"]:hover {
    background-color: #444;
    color: white;
    transition: all 0.3s ease;
}
</style>
"""
st.markdown(custom_css, unsafe_allow_html=True)

# Initialize or reset session state to ensure proper structure
if 'reset_session' not in st.session_state or st.session_state.reset_session:
    st.session_state.documents = {}
    st.session_state.selected_document_id = "all"
    st.session_state.chat_history = []
    st.session_state.combined_chunks = []
    st.session_state.combined_embeddings = None
    st.session_state.active_tab = "Chat"
    st.session_state.reset_session = False

# Make sure any individual variables are initialized
if "documents" not in st.session_state:
    st.session_state.documents = {}
if "selected_document_id" not in st.session_state:
    st.session_state.selected_document_id = "all"
if "chat_history" not in st.session_state:
    st.session_state.chat_history = []
if "combined_chunks" not in st.session_state:
    st.session_state.combined_chunks = []
if "combined_embeddings" not in st.session_state:
    st.session_state.combined_embeddings = None
if "active_tab" not in st.session_state:
    st.session_state.active_tab = "Chat"

# Add this code near the top where we initialize session state variables
if "reset_chat_input" not in st.session_state:
    st.session_state.reset_chat_input = False

# Define helper functions
def chunk_text(text: str, chunk_size: int = 1000, overlap: int = 200) -> List[str]:
    """Split text into overlapping chunks for better processing."""
    return split_into_chunks(text, chunk_size, overlap)

def display_metadata(metadata, filename):
    """Display comprehensive document metadata in a structured format."""
    try:
        # Two-column layout for better organization
        col1, col2 = st.columns(2)
        
        # Left column - Basic document info
        with col1:
            # Basic document information
            if "title" in metadata and metadata["title"]:
                st.markdown(f"**Title:** {metadata['title']}")
            
            if "contract_type" in metadata and metadata["contract_type"]:
                if isinstance(metadata["contract_type"], str):
                    st.markdown(f"**Type:** {metadata['contract_type']}")
            elif "document_type" in metadata and metadata["document_type"]:
                if isinstance(metadata["document_type"], str):
                    st.markdown(f"**Type:** {metadata['document_type']}")
            
            if "estimated_page_count" in metadata:
                st.markdown(f"**Estimated Pages:** {metadata['estimated_page_count']}")
            
            # Contract status
            if "contract_status" in metadata and metadata["contract_status"]:
                if isinstance(metadata["contract_status"], str):
                    status = metadata["contract_status"]
                    # Color the status based on its value
                    color = "green" if status.lower() in ["active", "executed"] else \
                            "red" if status.lower() in ["expired", "terminated"] else \
                            "orange" if status.lower() in ["draft", "under negotiation"] else "gray"
                    st.markdown(f"**Status:** <span style='color:{color}'>{status}</span>", unsafe_allow_html=True)
            
            # Version information
            if "version" in metadata and metadata["version"]:
                if isinstance(metadata["version"], str):
                    st.markdown(f"**Version:** {metadata['version']}")
            
            # Important dates section
            st.markdown("**Important Dates:**")
            date_fields = []
            
            # Check for specifically extracted date fields
            for date_field in ["effective_date", "execution_date", "termination_date"]:
                if date_field in metadata and metadata[date_field]:
                    if isinstance(metadata[date_field], str):
                        date_fields.append(f"{date_field.replace('_', ' ').title()}: {metadata[date_field]}")
            
            # If no specific date fields, fall back to generic dates
            if not date_fields and "dates" in metadata and metadata["dates"]:
                if isinstance(metadata["dates"], list):
                    date_fields = [f"• {date}" for date in metadata["dates"] if isinstance(date, str)]
                elif isinstance(metadata["dates"], str):
                    date_fields = [f"• {metadata['dates']}"]
            
            # Display dates
            for date in date_fields:
                st.markdown(f"• {date}")
        
        # Right column - Parties and organizations
        with col2:
            display_party_info(metadata)
        
        # Additional detailed metadata
        st.markdown("**Additional Contract Terms**")
        display_additional_fields(metadata)
        
        # Source document information
        if "source_document" in metadata and metadata["source_document"]:
            st.markdown("**Source Document Details**")
            source = metadata["source_document"]
            
            if isinstance(source, dict):
                for key, value in source.items():
                    if isinstance(value, (str, int, float, bool)):
                        st.markdown(f"**{key.replace('_', ' ').title()}:** {value}")
            elif isinstance(source, str):
                st.markdown(f"**Source:** {source}")
    
    except Exception as e:
        logger.error(f"Error displaying metadata: {str(e)}")
        st.error(f"Error displaying document metadata")

def display_party_info(metadata):
    """Display party and organization information safely."""
    try:
        # Parties/Organizations
        if "parties" in metadata and metadata["parties"]:
            st.markdown("**Parties:**")
            display_list_or_dict(metadata["parties"])
        elif "organizations" in metadata and metadata["organizations"]:
            st.markdown("**Organizations:**")
            display_list_or_dict(metadata["organizations"])
        else:
            st.markdown("• No organizations found")
        
        # People mentioned
        if "people" in metadata and metadata["people"]:
            st.markdown("**People:**")
            display_list_or_dict(metadata["people"])
        
        # Monetary values
        if "monetary_values" in metadata and metadata["monetary_values"]:
            st.markdown("**Monetary Values:**")
            display_list_or_dict(metadata["monetary_values"])
    
    except Exception as e:
        logger.error(f"Error displaying party info: {str(e)}")
        st.markdown("• Error displaying party information")

def display_list_or_dict(items):
    """Safely display a list or dictionary of items."""
    try:
        if isinstance(items, list):
            for item in items:
                if isinstance(item, dict) and "name" in item:
                    role = item.get("role", "")
                    if role:
                        st.markdown(f"• {item['name']} ({role})")
                    else:
                        st.markdown(f"• {item['name']}")
                elif isinstance(item, (str, int, float)):
                    st.markdown(f"• {item}")
                else:
                    st.markdown(f"• {str(item)}")
        elif isinstance(items, dict):
            for key, value in items.items():
                if isinstance(value, (str, int, float)):
                    st.markdown(f"• {key}: {value}")
                else:
                    st.markdown(f"• {key}: {str(value)}")
        elif isinstance(items, (str, int, float)):
            st.markdown(f"• {items}")
        else:
            st.markdown(f"• {str(items)}")
    
    except Exception as e:
        logger.error(f"Error displaying list or dict: {str(e)}")
        st.markdown("• Error displaying information")

def display_additional_fields(metadata):
    """Display additional metadata fields safely."""
    try:
        # Additional fields from the comprehensive list
        additional_fields = [
            # Legal Framework
            ("Jurisdiction", "jurisdiction"),
            ("Governing Law", "governing_law"),
            ("Contract Status", "contract_status"),
            ("Previous Version", "previous_version_reference"),
            ("Termination Notice", "termination_notice_period"),
            ("Automatic Renewal", "automatic_renewal"),
            ("Grounds for Termination", "grounds_for_termination"),
            
            # Commercial Terms
            ("Exclusivity", "exclusivity"),
            ("Exclusivity Agreement", "exclusivity_agreement"),
            ("Market Restrictions", "market_restrictions"),
            ("Non-Compete", "non_compete"),
            ("Competitor Restrictions", "competitor_restriction"),
            ("Post-Contract Restrictions", "post_contract_restriction_period"),
            
            # IP & Branding
            ("IP Assignment", "ip_assignment"),
            ("Trademark Ownership", "trademark_ownership"),
            ("Trademark Registration", "trademark_registration_status"),
            ("Trademark Usage License", "trademark_usage_license"),
            ("Branding Rights", "branding_rights"),
            ("Co-Branding Agreements", "co_branding_agreements"),
            
            # Financial Terms
            ("Payment Terms", "payment_obligations"),
            ("Royalty Percentage", "royalty_fee_percentage"),
            ("Revenue Share Model", "revenue_share_model"),
            ("Late Payment Penalty", "late_payment_penalty"),
            ("Revenue Collection", "revenue_collection_agent"),
            ("Performance Bonuses", "performance_bonuses"),
            
            # Operational Requirements
            ("Key Obligations", "key_obligations"),
            ("Service Standards", "service_standards"),
            ("Product Quality Standards", "product_quality_standards"),
            ("Compliance Requirements", "compliance_requirements"),
            ("Reporting Requirements", "reporting_requirements"),
            ("KPI Tracking", "kpi_tracking"),
            ("Audit Rights", "inspection_rights"),
            
            # Confidentiality & Data
            ("Confidentiality", "confidentiality_clause"),
            ("Confidentiality Duration", "confidentiality_duration"),
            ("Data Processing Agreement", "data_processing_agreement"),
            ("Third Party Disclosure", "third_party_disclosure_restrictions"),
            ("Sensitive Data Definition", "sensitive_data_definition"),
            ("Security Measures", "security_measures"),
            
            # Dispute & Breach
            ("Dispute Resolution", "dispute_resolution"),
            ("Force Majeure", "force_majeure"),
            ("Penalties for Breach", "penalties_for_breach"),
            ("Exit Compensation", "exit_compensation"),
            
            # Special Relationships
            ("Brand Licensor", "brand_licensor"),
            ("Licensee", "licensee"),
            ("Producers", "producers"),
            ("Partner Restaurants", "partner_restaurants"),
            ("Sub-Licensee", "sub_licensee"),
            
            # Marketing & Advertising
            ("Advertising Restrictions", "advertising_restrictions"),
            ("Marketing Approval", "marketing_approval_requirement"),
            ("Trademark in Ads", "use_of_trademark_in_ads"),
            ("Sales Channel Limitations", "sales_channel_limitations"),
            ("Influencer Restrictions", "influencer_advertising_restrictions")
        ]
        
        # Two columns for compact display
        col1, col2 = st.columns(2)
        
        for i, (display_name, field_name) in enumerate(additional_fields):
            # Alternate between columns
            with col1 if i % 2 == 0 else col2:
                if field_name in metadata and metadata[field_name]:
                    # Safely display the field value
                    if isinstance(metadata[field_name], (str, int, float, bool)):
                        st.markdown(f"**{display_name}:** {metadata[field_name]}")
                    else:
                        st.markdown(f"**{display_name}:** {str(metadata[field_name])}")
    
    except Exception as e:
        logger.error(f"Error displaying additional fields: {str(e)}")
        st.markdown("**Error displaying additional contract terms**")

def display_documents():
    """Display all documents with their metadata in expandable sections."""
    st.header("Document Library")
    
    if not st.session_state.documents:
        st.info("No documents uploaded yet. Use the sidebar to upload documents.")
        return
    
    for doc_id, doc in st.session_state.documents.items():
        try:
            # Handle different document formats safely
            if isinstance(doc, dict) and "filename" in doc:
                filename = doc["filename"]
                file_size = doc.get("file_size", 0)  # Use get with default for safety
                
                # Create expander with file info
                with st.expander(f"{filename} ({format_bytes(file_size) if file_size else 'Unknown size'})"):
                    # Display document metadata
                    if "metadata" in doc and doc["metadata"]:
                        display_metadata(doc["metadata"], filename)
                    else:
                        st.info("No metadata available for this document.")
            elif isinstance(doc, tuple):
                # Handle tuple case - extract what we can
                if len(doc) >= 1:
                    filename = doc_id  # Use the doc_id as the name
                    
                    with st.expander(f"{filename}"):
                        if len(doc) >= 2 and doc[1] is not None:
                            # Try to display metadata if it's in the tuple
                            display_metadata(doc[1], filename) 
                        else:
                            st.info("No metadata available for this document.")
            else:
                # Unknown format
                st.warning(f"Unknown document format for {doc_id}")
        except Exception as e:
            st.error(f"Error displaying document {doc_id}: {str(e)}")
            logging.error(f"Error displaying document {doc_id}: {str(e)}")

def initialize_embedding_model():
    """
    Initialize and store the sentence transformer model in session state.
    Call this during app startup to ensure model is available.
    """
    try:
        if "sentence_transformer" not in st.session_state or st.session_state.sentence_transformer is None:
            # Initialize the model
            model_name = "all-MiniLM-L6-v2"  # Fast and effective model
            logger.info(f"Loading SentenceTransformer model: {model_name}")
            
            # Set device appropriately (use GPU if available)
            device = "cuda" if torch.cuda.is_available() else "cpu"
            st.session_state.sentence_transformer = SentenceTransformer(model_name, device=device)
            logger.info(f"Successfully loaded SentenceTransformer using {device}")
            return True
        return True
    except Exception as e:
        logger.error(f"Error initializing embedding model: {str(e)}")
        return False

def create_embeddings(texts):
    """
    Create embeddings for the given texts using the sentence transformer model.
    """
    try:
        # Ensure model is initialized
        if not initialize_embedding_model():
            logger.error("Failed to initialize embedding model")
            return None
            
        model = st.session_state.sentence_transformer
        embeddings = model.encode(texts, convert_to_tensor=False, show_progress_bar=True)
        return embeddings
    except Exception as e:
        logger.error(f"Error creating embeddings: {str(e)}")
        return None

def initialize_faiss_index(dimension=384):
    """Initialize or reset the FAISS index with the given dimension."""
    try:
        # Create a new index
        logger.info(f"Initializing FAISS index with dimension {dimension}")
        index = faiss.IndexFlatL2(dimension)  # L2 distance
        st.session_state.faiss_index = index
        st.session_state.chunk_map = {}  # Map to track indices to chunks
        st.session_state.current_index = 0  # Track current index position
        logger.info("Successfully created new FAISS index")
        return True
    except Exception as e:
        logger.error(f"Error initializing FAISS index: {str(e)}")
        return False

def add_to_faiss_index(embeddings, chunks):
    """Add embeddings to the FAISS index and map them to chunks."""
    try:
        # Create/get index
        if "faiss_index" not in st.session_state or st.session_state.faiss_index is None:
            if not initialize_faiss_index(embeddings.shape[1]):
                return False
                
        # Initialize chunk map if not present
        if "chunk_map" not in st.session_state:
            st.session_state.chunk_map = {}
            
        # Initialize current index if not present
        if "current_index" not in st.session_state:
            st.session_state.current_index = 0
            
        # Get current index position
        start_idx = st.session_state.current_index
        
        # Add embeddings to the index
        st.session_state.faiss_index.add(np.array(embeddings).astype('float32'))
        
        # Map indices to chunks
        for i, chunk in enumerate(chunks):
            idx = start_idx + i
            st.session_state.chunk_map[idx] = chunk
            
        # Update current index
        st.session_state.current_index += len(chunks)
        
        logger.info(f"Added {len(chunks)} chunks to FAISS index. Total chunks: {st.session_state.current_index}")
        return True
    except Exception as e:
        logger.error(f"Error adding to FAISS index: {str(e)}")
        return False

def upload_and_process_document():
    """Upload and process a document with proper vector indexing."""
    uploaded_file = st.file_uploader("Upload a document", type=["pdf", "txt", "docx"])
    if uploaded_file is not None:
        with st.spinner("Processing document..."):
            # Process the document
            doc_data = process_document(uploaded_file)
            if doc_data:
                # Store document data
                if "documents" not in st.session_state:
                    st.session_state.documents = {}
                st.session_state.documents[uploaded_file.name] = doc_data
                
                # Extract chunks for vector indexing
                chunks = doc_data.get("chunks", [])
                
                # Create embeddings for chunks
                if chunks:
                    # Extract text from chunk dictionaries
                    texts = [chunk["text"] for chunk in chunks]
                    
                    # Create embeddings
                    embeddings = create_embeddings(texts)
                    
                    if embeddings is not None and len(embeddings) > 0:
                        # Add to FAISS index
                        if add_to_faiss_index(embeddings, chunks):
                            logger.info(f"Successfully indexed {len(chunks)} chunks from {uploaded_file.name}")
                            
                            # Store chunks in session state for retrieval
                            if "chunks" not in st.session_state:
                                st.session_state.chunks = []
                            st.session_state.chunks.extend(chunks)
                        else:
                            logger.error(f"Failed to index chunks from {uploaded_file.name}")
                    else:
                        logger.error("Failed to create embeddings for chunks")
                
                st.success(f"Successfully processed {uploaded_file.name}")
                return True
    return False

def search_similar(query, top_k=8):
    """
    Search for similar chunks in the FAISS index with improved chunk mapping.
    """
    try:
        # Ensure model is initialized
        if not initialize_embedding_model():
            logger.error("Failed to initialize embedding model for search")
            return []
        
        # Generate embedding directly using the model
        model = st.session_state.sentence_transformer
        embedding = model.encode([query], convert_to_tensor=False)
        
        # Check if FAISS index exists and has data
        if ("faiss_index" in st.session_state and 
            st.session_state.faiss_index is not None and 
            st.session_state.faiss_index.ntotal > 0):
            
            distances, indices = st.session_state.faiss_index.search(np.array(embedding).astype('float32'), top_k)
            
            # Format results using the chunk map
            results = []
            for i, (idx, dist) in enumerate(zip(indices[0], distances[0])):
                if idx == -1:  # FAISS returns -1 for empty slots
                    continue
                    
                # Use chunk map to get original chunk
                if "chunk_map" in st.session_state and idx in st.session_state.chunk_map:
                    chunk = st.session_state.chunk_map[idx]
                    results.append({
                        "index": int(idx),
                        "score": float(1.0 - dist),
                        "chunk": chunk
                    })
            
            logger.info(f"Found {len(results)} similar chunks")
            return results
        else:
            # This is important - log the reason for no results
            if "faiss_index" not in st.session_state or st.session_state.faiss_index is None:
                logger.warning("FAISS index not found. Returning empty results.")
            elif st.session_state.faiss_index.ntotal == 0:
                logger.warning("FAISS index is empty. Returning empty results.")
            return []
    except Exception as e:
        logger.error(f"Error in search_similar: {str(e)}")
        return []

def process_document(uploaded_file):
    """Process an uploaded document with enhanced context generation."""
    try:
        # Get file extension
        file_extension = os.path.splitext(uploaded_file.name)[1].lower()
        
        # Extract text based on file type
        if file_extension == ".pdf":
            text = extract_text_from_pdf(uploaded_file)
        elif file_extension in [".txt", ".md", ".html"]:
            text = uploaded_file.read().decode("utf-8")
        elif file_extension in [".docx", ".doc"]:
            text = extract_text_from_docx(uploaded_file)
        else:
            st.error(f"Unsupported file format: {file_extension}")
            return None
        
        # Extract metadata with the comprehensive extractor
        metadata_extractor = MetadataExtractor()
        metadata = metadata_extractor.extract_metadata(text, filename=uploaded_file.name)
        logger.info("Extracted metadata from document")
        
        # Split text into chunks for embedding
        chunks = split_text_into_chunks(text, source=uploaded_file.name)
        
        # Create document data structure
        doc_data = {
            "text": text,
            "metadata": metadata,
            "chunks": chunks,
            "filename": uploaded_file.name,
            "file_size": uploaded_file.size,
            # Don't generate summary here - wait for user to request it
        }
        
        return doc_data
        
    except Exception as e:
        logger.error(f"Error processing document: {str(e)}")
        st.error(f"Error processing document: {str(e)}")
        return None

def process_uploaded_files():
    """Process all uploaded files and store their data in session state."""
    if 'documents' not in st.session_state:
        st.session_state.documents = {}
    
    # Get file uploader from the sidebar with better instructions
    with st.sidebar:
        st.info("📌 Upload PDF documents to analyze, summarize, and chat with them using AI.")
        
        uploaded_files = st.file_uploader(
            "Upload PDF Documents",
            type=["pdf"],
            accept_multiple_files=True,
            key="file_uploader",
            help="Drag and drop PDF files here or click to select files"
        )
        
        if not uploaded_files:
            st.warning("No documents uploaded yet. Please upload PDF files to begin analysis.")
    
    if uploaded_files:
        # Process each uploaded file
        for uploaded_file in uploaded_files:
            # Check if we've already processed this file
            file_id = uploaded_file.name
            
            if file_id not in st.session_state.documents:
                with st.spinner(f"Processing {uploaded_file.name}..."):
                    # Save the file and extract text
                    temp_path = save_uploaded_file(uploaded_file)
                    text = extract_text_from_pdf(uploaded_file)
                    
                    if text:
                        # Extract metadata
                        metadata = extract_metadata(text, uploaded_file.name)
                        
                        # Store the document data
                        st.session_state.documents[file_id] = {
                            "filename": uploaded_file.name,
                            "text": text,
                            "metadata": metadata,
                            "file_size": uploaded_file.size,
                            "chunks": [],
                            "temp_path": temp_path
                        }
                        
                        # Create chunks and embeddings for RAG
                        chunks = chunk_text(text)
                        if chunks:
                            embeddings = create_embeddings(chunks)
                            
                            # Only store if embeddings were created successfully
                            if embeddings is not None:
                                st.session_state.documents[file_id]["chunks"] = chunks
                                st.session_state.documents[file_id]["embeddings"] = embeddings
                                st.session_state.documents[file_id]["index"] = create_index(embeddings)
                    else:
                        st.error(f"Could not extract text from {uploaded_file.name}")
        
        # Show success message once all files are processed
        st.sidebar.success(f"Processed {len(uploaded_files)} document(s)")

def get_document_context(query, max_chunks=8):
    """
    Enhanced document context retrieval with direct chunk access.
    """
    try:
        # Get document ID if specific document is selected
        doc_id = st.session_state.selected_document_id
        
        # Initialize context parts
        context_parts = []
        
        # Get search results with error handling
        search_results = search_similar(query, top_k=max_chunks)
        
        if search_results:
            # Extract the most relevant chunks directly from the search results
            chunks_section = "MOST RELEVANT TEXT PASSAGES:\n"
            for result in search_results:
                chunk = result.get("chunk", {})
                source = chunk.get("source", "Unknown")
                text = chunk.get("text", "")
                if text:
                    chunks_section += f"\n[From: {source}]\n{text}\n"
            
            context_parts.append(chunks_section)
            logger.info(f"Added {len(search_results)} relevant chunks to context")
        else:
            logger.warning("No relevant chunks found from vector search")
            
            # Fallback: Add sample chunks if no search results
            sample_chunks_section = "DOCUMENT SAMPLES:\n"
            sample_added = False
            
            # Include sample text from documents
            if doc_id == "all":
                # Add samples from all documents (up to 3)
                for doc_name, doc_data in list(st.session_state.documents.items())[:3]:
                    text = doc_data.get("text", "")
                    if text:
                        # Take first 500 characters as sample
                        sample = text[:500] + "..." if len(text) > 500 else text
                        sample_chunks_section += f"\n[Sample from: {doc_name}]\n{sample}\n"
                        sample_added = True
            else:
                # Add sample from specific document
                if doc_id in st.session_state.documents:
                    text = st.session_state.documents[doc_id].get("text", "")
                    if text:
                        # Take first 1000 characters as sample for specific document
                        sample = text[:1000] + "..." if len(text) > 1000 else text
                        sample_chunks_section += f"\n[Sample from: {doc_id}]\n{sample}\n"
                        sample_added = True
            
            # Only add sample section if we found samples
            if sample_added:
                context_parts.append(sample_chunks_section)
                logger.info("Added document samples as fallback")
        
        # Add document summaries
        summary_section = "DOCUMENT SUMMARIES:\n"
        summary_added = False
        
        if doc_id == "all":
            # Add summaries from all documents
            for doc_name, doc_data in st.session_state.documents.items():
                if isinstance(doc_data, dict) and "summary" in doc_data and doc_data["summary"]:
                    summary_section += f"\n[Summary of {doc_name}]\n{doc_data['summary']}\n"
                    summary_added = True
        else:
            # Add summary for specific document
            if doc_id in st.session_state.documents:
                doc_data = st.session_state.documents[doc_id]
                if isinstance(doc_data, dict) and "summary" in doc_data and doc_data["summary"]:
                    summary_section += f"\n[Summary of {doc_id}]\n{doc_data['summary']}\n"
                    summary_added = True
        
        # Only add summary section if we found summaries
        if summary_added:
            context_parts.append(summary_section)
            logger.info("Added document summaries to context")
        
        # Add document metadata
        metadata_section = "DOCUMENT METADATA:\n"
        metadata_added = False
        
        if doc_id == "all":
            # Add metadata from all documents
            for doc_name, doc_data in st.session_state.documents.items():
                if isinstance(doc_data, dict) and "metadata" in doc_data and doc_data["metadata"]:
                    metadata = doc_data["metadata"]
                    
                    # Extract key metadata fields
                    meta_text = f"\n[Metadata for {doc_name}]\n"
                    
                    # Add important metadata fields
                    for key, value in metadata.items():
                        if value and key not in ["text", "content"]:  # Skip large text fields
                            meta_text += f"{key}: {value}\n"
                    
                    # Always include filename which often contains version/date info
                    meta_text += f"filename: {doc_name}\n"
                    
                    # Parse dates from filename for version-related queries
                    date_match = re.search(r'(\d{2,4})[-_\s]?(\d{1,2})[-_\s]?(\d{1,2})', doc_name)
                    if date_match:
                        year, month, day = date_match.groups()
                        if len(year) == 2:
                            year = f"20{year}"  # Assume 20xx for 2-digit years
                        meta_text += f"filename_date: {year}-{month.zfill(2)}-{day.zfill(2)}\n"
                    
                    metadata_section += meta_text
                    metadata_added = True
                else:
                    # Even if no metadata, include filename
                    metadata_section += f"\n[Document: {doc_name}]\n"
                    metadata_section += f"filename: {doc_name}\n"
                    metadata_added = True
        else:
            # Include detailed metadata for the specific document
            if doc_id in st.session_state.documents:
                doc_data = st.session_state.documents[doc_id]
                if isinstance(doc_data, dict) and "metadata" in doc_data and doc_data["metadata"]:
                    metadata = doc_data["metadata"]
                    meta_text = f"\n[Metadata for {doc_id}]\n"
                    
                    # Include all metadata fields for single document view
                    for key, value in metadata.items():
                        if value and key not in ["text", "content"]:  # Skip large text fields
                            meta_text += f"{key}: {value}\n"
                    
                    # Always include filename
                    meta_text += f"filename: {doc_id}\n"
                    
                    metadata_section += meta_text
                    metadata_added = True
                else:
                    # Even if no metadata, include filename
                    metadata_section += f"\n[Document: {doc_id}]\n"
                    metadata_section += f"filename: {doc_id}\n"
                    metadata_added = True
        
        # Only add metadata section if we found metadata
        if metadata_added:
            context_parts.append(metadata_section)
            logger.info("Added document metadata to context")
        
        # Add document filenames section
        filenames_section = "DOCUMENT FILENAMES:\n"
        if doc_id == "all":
            for doc_name in st.session_state.documents.keys():
                filenames_section += f"- {doc_name}\n"
            context_parts.append(filenames_section)
            logger.info("Added document filenames to context")
        else:
            filenames_section += f"- {doc_id}\n"
            context_parts.append(filenames_section)
            logger.info("Added document filename to context")
        
        # Combine the context parts with clear section dividers
        if context_parts:
            context = "\n\n" + "\n\n".join(context_parts) + "\n\n"
            logger.info(f"Created complete context with {len(context_parts)} sections")
            return context
        else:
            return "No relevant information found in the document(s)."
            
    except Exception as e:
        logger.error(f"Error getting document context: {str(e)}")
        # Emergency fallback - at least provide document names
        emergency_context = "DOCUMENT FILENAMES:\n"
        for doc_name in st.session_state.documents.keys():
            emergency_context += f"- {doc_name}\n"
        return emergency_context
        
        # The rest of the function (metadata, filenames, etc.) can stay the same
        # ...

def process_user_query(query, document_context):
    """Process user query with enhanced context instructions."""
    try:
        # Craft a comprehensive prompt that uses all context sources
        prompt = f"""You are a precise legal document assistant. Answer questions directly using ALL the provided document information.

DOCUMENT INFORMATION:
{document_context}

USER QUESTION: {query}

ANSWER GUIDELINES:
1. Use ALL available information sources provided:
   - DOCUMENT TEXT (relevant passages from the documents)
   - DOCUMENT SUMMARIES (for high-level overview)
   - DOCUMENT METADATA (for specific details about versions, dates, parties, etc.)
   - DOCUMENT FILENAMES (which may contain version/date information)

2. For questions about versions or dates:
   - Check filenames for date patterns (YYYY-MM-DD)
   - Compare dates to determine the most recent document
   - Return the EXACT FILENAME of the most recent document and its date

3. For questions about specific clauses:
   - First check relevant text passages
   - Reference the document summary for context
   - Cite specific details from the metadata

4. Keep your answer brief but comprehensive
5. Provide specific facts, not general statements
6. If information isn't available, clearly state what's missing

DIRECT ANSWER:"""

        # Get response from LLM with low temperature for factual accuracy
        response = get_llama_response(prompt, temperature=0.1, max_tokens=300)
        
        # Clean up common prefacing text
        cleaned_response = re.sub(r'^(Based on|According to|From the|The document|In the|As per)', '', response, flags=re.IGNORECASE).strip()
        
        return cleaned_response if cleaned_response else response
    except Exception as e:
        logger.error(f"Error processing query: {str(e)}")
        return f"Error processing your question: {str(e)}"

# Replace the existing chat handling with this simpler function
def handle_chat():
    """Handle the chat interface with improved error handling and direct answers."""
    st.header("Chat with Documents")
    
    # Document selector
    st.write("Select a document to chat with:")
    doc_options = ["All Documents"] + list(st.session_state.documents.keys())
    selected_doc = st.selectbox(
        "Select a document",
        options=doc_options,
        index=0,
        key="document_selector",
        label_visibility="collapsed"
    )
    
    # Update session state
    st.session_state.selected_document_id = selected_doc if selected_doc != "All Documents" else "all"
    
    # Initialize chat history if needed
    if "chat_history" not in st.session_state:
        st.session_state.chat_history = []
    
    # Display chat history
    for message in st.session_state.chat_history:
        st.markdown(f"**You:** {message['user']}")
        st.markdown(f"**AI:** {message['ai']}")
        st.markdown("---")
    
    # Define callback function for form submission
    def submit_query():
        if st.session_state.temp_query and st.session_state.temp_query.strip():
            # Store the query value
            query = st.session_state.temp_query
            
            # Add to chat history immediately (before processing)
            st.session_state.chat_history.append({
                "user": query,
                "ai": "Processing your question..."
            })
            
            # Set flag to process this query after rerun
            st.session_state.process_query = query
            
            # Clear the input field by setting temp_query to empty
            st.session_state.temp_query = ""
    
    # Create a form for input
    with st.form(key="query_form", clear_on_submit=True):
        # Initialize temp_query if not exists
        if "temp_query" not in st.session_state:
            st.session_state.temp_query = ""
            
        # Use text input with key temp_query
        st.text_input(
            "Ask a question about your document(s):",
            key="temp_query", 
            placeholder="e.g., What is the last contract version?"
        )
        
        # Submit button
        submit_button = st.form_submit_button("Send", on_click=submit_query)
    
    # Process query after form submission (and page rerun)
    if "process_query" in st.session_state and st.session_state.process_query:
        query = st.session_state.process_query
        
        try:
            # Get document context
            doc_context = get_document_context(query)
            
            # Process the query
            answer = process_user_query(query, doc_context)
            
            # Update the last entry in chat history
            st.session_state.chat_history[-1]["ai"] = answer
            
        except Exception as e:
            error_msg = f"Sorry, I encountered an error: {str(e)}"
            st.session_state.chat_history[-1]["ai"] = error_msg
            logger.error(f"Chat error: {str(e)}")
        
        # Clear the process flag
        st.session_state.process_query = None
        
        # Force a rerun to show the updated chat
        st.rerun()

def find_latest_contract(documents: dict) -> str:
    """
    Analyze document metadata to identify the latest contract version by date.
    """
    try:
        # Extract dates from all documents
        contract_dates = []
        
        for doc_id, doc_data in documents.items():
            logger.info(f"Examining document dates for: {doc_id}")
            
            # Extract date from filename if present (often more reliable)
            date_from_filename = None
            filename = doc_data.get("filename", doc_id)
            
            # Look for date in format like "03 02 2020" or similar in filename
            date_patterns = [
                r'(\d{2}\s*\d{2}\s*\d{4})',  # DD MM YYYY with optional spaces
                r'(\d{4}\s*\d{2}\s*\d{2})'   # YYYY MM DD with optional spaces
            ]
            
            for pattern in date_patterns:
                matches = re.findall(pattern, filename)
                if matches:
                    date_str = matches[0].replace(" ", "")
                    # Try to standardize to YYYY-MM-DD
                    if len(date_str) == 8:  # 8 digits like 03022020 or 20200302
                        if int(date_str[:2]) <= 31 and int(date_str[2:4]) <= 12:
                            # Likely DD MM YYYY format
                            date_from_filename = f"20{date_str[4:6]}-{date_str[2:4]}-{date_str[:2]}"
                        else:
                            # Likely YYYY MM DD format
                            date_from_filename = f"{date_str[:4]}-{date_str[4:6]}-{date_str[6:8]}"
                    logger.info(f"Extracted date from filename: {date_from_filename}")
            
            # Try to get date from metadata
            if "metadata" in doc_data and doc_data["metadata"]:
                meta = doc_data["metadata"]
                
                # Look for various date fields
                date_fields = ["effective_date", "execution_date", "date"]
                
                for field in date_fields:
                    if field in meta and meta[field]:
                        date_value = meta[field]
                        logger.info(f"Found date in metadata field '{field}': {date_value}")
                        
                        # Add to our list with document ID
                        contract_dates.append({
                            "doc_id": doc_id,
                            "filename": filename,
                            "date": date_value,
                            "source_field": field
                        })
            
            # If we found a date in the filename, use it as well
            if date_from_filename:
                contract_dates.append({
                    "doc_id": doc_id,
                    "filename": filename,
                    "date": date_from_filename,
                    "source_field": "filename"
                })
        
        # Log all found dates for debugging
        logger.info(f"Found dates: {contract_dates}")
        
        # If we found dates, sort them newest first using proper date comparison
        if contract_dates:
            # Sort by date using a parsing function that handles various formats
            def get_sortable_date(date_obj):
                date_str = date_obj["date"]
                # Convert to YYYY-MM-DD if possible for proper sorting
                try:
                    # Handle different date formats
                    if re.match(r'\d{4}-\d{2}-\d{2}', date_str):
                        return date_str  # Already in YYYY-MM-DD
                    elif re.match(r'\d{2}/\d{2}/\d{4}', date_str):
                        parts = date_str.split('/')
                        return f"{parts[2]}-{parts[0].zfill(2)}-{parts[1].zfill(2)}"
                    elif re.match(r'\d{2}-\d{2}-\d{4}', date_str):
                        parts = date_str.split('-')
                        return f"{parts[2]}-{parts[0].zfill(2)}-{parts[1].zfill(2)}"
                    else:
                        return date_str
                except:
                    return date_str  # Return original if parsing fails
            
            # Sort contracts by date (newest first)
            sorted_contracts = sorted(contract_dates, key=lambda x: get_sortable_date(x), reverse=True)
            
            # Log sorted results for debugging
            logger.info(f"Sorted dates: {[c['date'] for c in sorted_contracts]}")
            
            # Return information about the latest contract
            latest = sorted_contracts[0]
            return f"""Based on date analysis, the latest contract version is:
- **Document:** {latest['filename']}
- **Date:** {latest['date']} (from {latest['source_field']})

All contracts with dates (sorted by newest first):
{chr(10).join([f"- {c['filename']}: {c['date']}" for c in sorted_contracts[:5]])}
"""
        else:
            return "No date information found in the available contracts."
    
    except Exception as e:
        logger.error(f"Error finding latest contract: {str(e)}")
        import traceback
        logger.error(traceback.format_exc())
        return f"Error analyzing contract dates: {str(e)}"

# Add this function for date debugging information
def debug_date_comparison(date1, date2):
    """Debug helper for date comparison issues"""
    logger.info(f"Comparing dates: '{date1}' vs '{date2}'")
    
    # Try to standardize dates
    try:
        # Handle various date formats
        def standardize_date(date_str):
            # Already in YYYY-MM-DD
            if re.match(r'\d{4}-\d{2}-\d{2}', date_str):
                return date_str
            # MM/DD/YYYY format
            elif re.match(r'\d{2}/\d{2}/\d{4}', date_str):
                parts = date_str.split('/')
                return f"{parts[2]}-{parts[0].zfill(2)}-{parts[1].zfill(2)}"
            # DD-MM-YYYY format
            elif re.match(r'\d{2}-\d{2}-\d{4}', date_str):
                parts = date_str.split('-')
                return f"{parts[2]}-{parts[0].zfill(2)}-{parts[1].zfill(2)}"
            # Other format
            else:
                return date_str
        
        std_date1 = standardize_date(date1)
        std_date2 = standardize_date(date2)
        
        logger.info(f"Standardized: '{std_date1}' vs '{std_date2}'")
        logger.info(f"Comparison result: {std_date1 > std_date2}")
        
        return std_date1, std_date2
    except Exception as e:
        logger.error(f"Error in date comparison: {str(e)}")
        return date1, date2

def split_text_into_chunks(text, chunk_size=1000, overlap=200, source="unknown"):
    """
    Split text into overlapping chunks for embedding.
    
    Args:
        text (str): The text to split
        chunk_size (int): Size of each chunk
        overlap (int): Overlap between chunks
        source (str): Source document name
        
    Returns:
        list: List of chunk objects with text and metadata
    """
    try:
        logger.info(f"Splitting text into chunks (size={chunk_size}, overlap={overlap})")
        
        # Handle empty or None text
        if not text:
            logger.warning("Empty text provided for chunking")
            return []
            
        # Split into sentences (simple approach - can be improved)
        sentences = re.split(r'(?<=[.!?])\s+', text)
        
        chunks = []
        current_chunk = ""
        current_size = 0
        
        for sentence in sentences:
            # Skip empty sentences
            if not sentence.strip():
                continue
                
            sentence_len = len(sentence)
            
            # If this sentence alone exceeds chunk size, break it up further
            if sentence_len > chunk_size:
                # If current chunk has content, save it
                if current_chunk:
                    chunks.append({
                        "text": current_chunk,
                        "source": source
                    })
                
                # Break long sentence into smaller pieces
                for i in range(0, sentence_len, chunk_size - overlap):
                    chunk_text = sentence[i:i + chunk_size]
                    if chunk_text:
                        chunks.append({
                            "text": chunk_text,
                            "source": source
                        })
                
                current_chunk = ""
                current_size = 0
                continue
            
            # If adding this sentence would exceed the chunk size, save current chunk and start a new one
            if current_size + sentence_len > chunk_size:
                if current_chunk:
                    chunks.append({
                        "text": current_chunk,
                        "source": source
                    })
                
                # Start new chunk with overlap from previous chunk if possible
                if overlap > 0 and current_size > overlap:
                    # Get the last part of the previous chunk for overlap
                    overlap_text = current_chunk[-overlap:]
                    current_chunk = overlap_text + " " + sentence
                    current_size = len(overlap_text) + 1 + sentence_len
                else:
                    current_chunk = sentence
                    current_size = sentence_len
            else:
                # Add sentence to current chunk
                if current_chunk:
                    current_chunk += " " + sentence
                    current_size += 1 + sentence_len
                else:
                    current_chunk = sentence
                    current_size = sentence_len
        
        # Add the last chunk if it has content
        if current_chunk:
            chunks.append({
                "text": current_chunk,
                "source": source
            })
        
        logger.info(f"Split text into {len(chunks)} chunks")
        return chunks
        
    except Exception as e:
        logger.error(f"Error splitting text into chunks: {str(e)}")
        return []

def extract_text_from_docx(docx_file):
    """
    Extract text from a DOCX file using python-docx library.
    
    Args:
        docx_file: The uploaded DOCX file
        
    Returns:
        str: Extracted text from the document
    """
    try:
        import io
        from docx import Document
        
        # Create a BytesIO object from the uploaded file
        bytes_io = io.BytesIO(docx_file.getvalue())
        
        # Open the document
        document = Document(bytes_io)
        
        # Extract text from paragraphs
        full_text = []
        for para in document.paragraphs:
            full_text.append(para.text)
            
        # Also get text from tables
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    full_text.append(cell.text)
        
        # Join all text parts with newlines
        return '\n'.join(full_text)
    except Exception as e:
        logger.error(f"Error extracting text from DOCX: {str(e)}")
        return ""

def extract_metadata(text, filename, max_length=2500):
    """
    Extract comprehensive metadata from document text using the LLM.
    
    Args:
        text (str): The document text
        filename (str): Name of the file
        max_length (int): Maximum text length to send to API
        
    Returns:
        dict: Detailed metadata extracted from the document
    """
    try:
        logger.info(f"Extracting comprehensive metadata from document: {filename}")
        
        # Truncate text if needed to avoid token limits
        if len(text) > max_length:
            truncated_text = text[:max_length] + "..."
            logger.info(f"Truncated text from {len(text)} to {len(truncated_text)} characters for metadata extraction")
        else:
            truncated_text = text
            
        # Create enhanced prompt with comprehensive field list
        prompt = (
            f"Extract metadata from the following legal document (filename: {filename}) as a valid JSON object. "
            "Include each field only if it is found in the text.\n\n"
            "==== METADATA EXTRACTION GUIDELINES ====\n"
            "Include the following fields **only if they appear** in the text:\n"
            "- contract_type: The type of legal document (e.g., NDA, Service Agreement, Distribution Agreement).\n"
            "- parties: The organizations or individuals involved.\n"
            "- effective_date: The date when the agreement takes effect (YYYY-MM-DD).\n"
            "- execution_date: The date when the document was signed.\n"
            "- termination_date: The date when the agreement ends or 'Indefinite'.\n"
            "- jurisdiction: The governing jurisdiction (e.g., France, EU, USA).\n"
            "- governing_law: The legal framework (e.g., French Law, EU Regulations).\n"
            "- version: The contract version or amendment indicator (e.g., V1, V2, Final, Draft).\n\n"
            "Additional Metadata (if present):\n"
            "- contract_status: Current status (Active, Expired, Terminated, Under Negotiation).\n"
            "- previous_version_reference: Reference to prior version(s).\n"
            "- key_obligations: Main responsibilities of the parties.\n"
            "- payment_obligations: Payment terms or financial obligations.\n"
            "- confidentiality_clause: Details of confidentiality or data protection obligations.\n"
            "- dispute_resolution: Mechanisms for resolving disputes.\n"
            "- force_majeure: Conditions excusing performance.\n"
            "- exclusivity: Whether one party has exclusive rights.\n"
            "- non_compete: Restrictions on engaging with competitors.\n"
            "- ip_assignment: Ownership rights or licensing.\n\n"
            "==== SPECIFIC CONTRACT FIELDS ====\n"
            "Also include these fields if they appear in the text:\n"
            "- brand_licensor: Entity granting rights over brand use.\n"
            "- licensee: Entity receiving brand/IP rights.\n"
            "- royalty_fee_percentage: Percentage payable as a royalty fee.\n"
            "- revenue_share_model: Details on revenue sharing between parties.\n"
            "- termination_notice_period: Notice period required for termination.\n"
            "- automatic_renewal: Whether the contract renews automatically.\n"
            "- grounds_for_termination: Conditions that allow contract termination.\n"
            "- exclusivity_agreement: Details on any exclusivity agreements.\n"
            "- market_restrictions: Geographic or sector-specific limitations.\n"
            "- confidentiality_duration: Duration for which confidentiality obligations persist.\n\n"
            "==== FILENAME & DATE DETECTION ====\n"
            "- If the filename contains a date pattern (e.g., '2023-01-05'), include it as 'parsed_date'.\n"
            "- Determine if this appears to be the latest version based on the date and version info.\n"
            "- Include this information under 'source_document'.\n\n"
            "==== DOCUMENT TEXT ====\n"
            f"{truncated_text}\n\n"
            "Return ONLY a valid JSON object with the detected fields. Omit any field not found in the text. "
            "Use the following structure:\n"
            "{\n"
            "  \"metadata\": { /* All extracted metadata fields */ },\n"
            "  \"source_document\": {\n"
            "    \"filename\": \"" + filename + "\",\n"
            "    \"parsed_date\": \"YYYY-MM-DD\" /* if date found in filename */,\n"
            "    \"is_latest_version\": true/false /* based on available information */\n"
            "  }\n"
            "}"
        )

        # Get LLM response with increased max tokens
        response = get_llama_response(prompt, temperature=0, max_tokens=1500)
        logger.info("Received metadata extraction response from LLM")
        
        # Find JSON in the response - look for anything between curly braces
        json_match = re.search(r'({[\s\S]*})', response)
        if json_match:
            json_str = json_match.group(1)
            
            # Try to parse the JSON
            try:
                result = json.loads(json_str)
                
                # Extract metadata from the result structure
                metadata = result.get("metadata", {})
                
                # Add filename to metadata if not already present
                if "filename" not in metadata:
                    metadata["filename"] = filename
                
                # Add source document info if available
                source_doc = result.get("source_document", {})
                if source_doc:
                    for key, value in source_doc.items():
                        metadata[f"source_{key}"] = value
                
                logger.info(f"Successfully extracted {len(metadata)} metadata fields")
                return metadata
                
            except json.JSONDecodeError as e:
                logger.error(f"Error parsing JSON from LLM response: {str(e)}")
                # Attempt to clean the JSON string
                clean_json_str = clean_json(json_str)
                
                try:
                    cleaned_result = json.loads(clean_json_str)
                    metadata = cleaned_result.get("metadata", {})
                    if "filename" not in metadata:
                        metadata["filename"] = filename
                    
                    # Add source document info if available
                    source_doc = cleaned_result.get("source_document", {})
                    if source_doc:
                        for key, value in source_doc.items():
                            metadata[f"source_{key}"] = value
                    
                    logger.info(f"Successfully extracted {len(metadata)} metadata fields after cleaning JSON")
                    return metadata
                except:
                    logger.error("Failed to parse JSON even after cleaning")
                    # Extract date from filename as fallback
                    date_match = re.search(r'(\d{4})[-_]?(\d{1,2})[-_]?(\d{1,2})', filename)
                    basic_metadata = {"filename": filename}
                    if date_match:
                        year, month, day = date_match.groups()
                        basic_metadata["source_parsed_date"] = f"{year}-{month.zfill(2)}-{day.zfill(2)}"
                    
                    return basic_metadata
        
        # Fallback if no JSON found
        logger.warning("No valid JSON found in LLM response for metadata extraction")
        # Extract date from filename as fallback
        date_match = re.search(r'(\d{4})[-_]?(\d{1,2})[-_]?(\d{1,2})', filename)
        basic_metadata = {"filename": filename}
        if date_match:
            year, month, day = date_match.groups()
            basic_metadata["source_parsed_date"] = f"{year}-{month.zfill(2)}-{day.zfill(2)}"
        
        return basic_metadata
        
    except Exception as e:
        logger.error(f"Error extracting metadata: {str(e)}")
        return {"filename": filename}  # Return at least the filename

def clean_json(json_str):
    """
    Attempt to clean and fix common JSON formatting issues.
    """
    # Replace single quotes with double quotes
    json_str = json_str.replace("'", '"')
    
    # Remove trailing commas in objects and arrays
    json_str = re.sub(r',\s*}', '}', json_str)
    json_str = re.sub(r',\s*]', ']', json_str)
    
    # Fix boolean and null values
    json_str = re.sub(r':\s*True', ': true', json_str)
    json_str = re.sub(r':\s*False', ': false', json_str)
    json_str = re.sub(r':\s*None', ': null', json_str)
    
    # Fix missing quotes around keys
    json_str = re.sub(r'([{,]\s*)([a-zA-Z0-9_]+)(\s*:)', r'\1"\2"\3', json_str)
    
    return json_str

# Main app layout
def main():
    """Main function for the Streamlit app with tab-based interface."""
    # Initialize session state variables
    if "chat_history" not in st.session_state:
        st.session_state.chat_history = []
    if "documents" not in st.session_state:
        st.session_state.documents = {}
    if "selected_document_id" not in st.session_state:
        st.session_state.selected_document_id = "all"
        
    # Initialize embedding model
    initialize_embedding_model()
    
    # Sidebar for document upload and selection
    with st.sidebar:
        st.header("Document Settings")
        
        # Document upload section
        st.subheader("Upload Documents")
        uploaded_files = st.file_uploader("Upload document(s)", 
                                         type=["pdf", "txt", "docx"], 
                                         accept_multiple_files=True,
                                         key="document_uploader")
        
        # Process uploaded files
        if uploaded_files:
            for uploaded_file in uploaded_files:
                # Process only new files
                if uploaded_file.name not in st.session_state.documents:
                    with st.spinner(f"Processing {uploaded_file.name}..."):
                        # Process the document
                        doc_data = process_document(uploaded_file)
                        if doc_data:
                            # Store document data
                            st.session_state.documents[uploaded_file.name] = doc_data
                            
                            # Create embeddings and add to FAISS index
                            chunks = doc_data.get("chunks", [])
                            if chunks:
                                texts = [chunk["text"] for chunk in chunks]
                                embeddings = create_embeddings(texts)
                                if embeddings is not None and len(embeddings) > 0:
                                    if add_to_faiss_index(embeddings, chunks):
                                        logger.info(f"Successfully indexed {len(chunks)} chunks from {uploaded_file.name}")
                                    
                            st.success(f"Processed {uploaded_file.name}")
        
        # Document selection
        st.subheader("Select Document")
        if st.session_state.documents:
            doc_options = list(st.session_state.documents.keys())
            doc_options.insert(0, "All Documents")
            selected_doc = st.selectbox("Choose a document", doc_options, key="doc_select")
            
            if selected_doc == "All Documents":
                st.session_state.selected_document_id = "all"
            else:
                st.session_state.selected_document_id = selected_doc
        else:
            st.info("No documents uploaded yet.")
    
    # Main area
    st.title("Document Chat")
    st.write("Upload documents and chat with them using AI")
    
    # Create tabs
    tabs = st.tabs(["💬 Chat History", "📄 Metadata", "📝 Summary", "🔄 Comparison"])
    
    has_documents = len(st.session_state.documents) > 0
    
    # Chat History tab
    with tabs[0]:
        if not has_documents:
            st.info("Please upload documents to start chatting.")
        else:
            # Show selected document
            if st.session_state.selected_document_id != "all":
                st.write(f"**Currently chatting with:** {st.session_state.selected_document_id}")
            else:
                st.write("**Currently chatting with:** All documents")
            
            # Display chat history
            for message in st.session_state.chat_history:
                with st.chat_message(message["role"]):
                    st.write(message["content"])
    
    # Metadata tab
    with tabs[1]:
        if not has_documents:
            st.info("Please upload documents to view metadata.")
        else:
            # Display metadata for selected document or all documents
            if st.session_state.selected_document_id != "all":
                # Show metadata for selected document
                doc_data = st.session_state.documents.get(st.session_state.selected_document_id, {})
                metadata = doc_data.get("metadata", {})
                
                if metadata:
                    st.subheader(f"Metadata for {st.session_state.selected_document_id}")
                    st.json(metadata)
                else:
                    st.info(f"No metadata available for {st.session_state.selected_document_id}")
            else:
                # Show metadata for all documents
                st.subheader("Metadata for All Documents")
                for doc_name, doc_data in st.session_state.documents.items():
                    with st.expander(f"Metadata for {doc_name}"):
                        metadata = doc_data.get("metadata", {})
                        if metadata:
                            st.json(metadata)
                        else:
                            st.info("No metadata available")
    
    # Summary tab
    with tabs[2]:
        if not has_documents:
            st.info("Please upload documents to generate summaries.")
        else:
            if st.session_state.selected_document_id != "all":
                # Generate summary for selected document
                doc_name = st.session_state.selected_document_id
                doc_data = st.session_state.documents.get(doc_name, {})
                
                if "summary" in doc_data:
                    st.subheader(f"Summary for {doc_name}")
                    st.markdown(doc_data["summary"])
                else:
                    if st.button(f"Generate Summary for {doc_name}"):
                        with st.spinner("Generating summary..."):
                            summary = generate_document_summary(doc_data.get("text", ""))
                            if summary:
                                doc_data["summary"] = summary
                                st.session_state.documents[doc_name] = doc_data
                                st.markdown(summary)
                                st.success("Summary generated!")
                            else:
                                st.error("Failed to generate summary.")
            else:
                # Option to generate summaries for all documents
                st.subheader("Document Summaries")
                for doc_name, doc_data in st.session_state.documents.items():
                    with st.expander(f"Summary for {doc_name}"):
                        if "summary" in doc_data:
                            st.markdown(doc_data["summary"])
                        else:
                            if st.button(f"Generate Summary for {doc_name}", key=f"sum_{doc_name}"):
                                with st.spinner("Generating summary..."):
                                    summary = generate_document_summary(doc_data.get("text", ""))
                                    if summary:
                                        doc_data["summary"] = summary
                                        st.session_state.documents[doc_name] = doc_data
                                        st.markdown(summary)
                                        st.success("Summary generated!")
                                    else:
                                        st.error("Failed to generate summary.")
    
    # Comparison tab
    with tabs[3]:
        if len(st.session_state.documents) < 2:
            st.info("Please upload at least two documents to compare them.")
        else:
            # Let user select two documents to compare
            doc_names = list(st.session_state.documents.keys())
            col1, col2 = st.columns(2)
            with col1:
                first_doc = st.selectbox("First Document", doc_names, key="first_doc")
            with col2:
                # Filter out the first document from options
                second_options = [doc for doc in doc_names if doc != first_doc]
                second_doc = st.selectbox("Second Document", second_options, key="second_doc")
            
            if st.button("Compare Documents", key="compare_btn"):
                with st.spinner("Comparing documents..."):
                    # Get the text content of both documents
                    doc1_text = st.session_state.documents[first_doc].get("text", "")
                    doc2_text = st.session_state.documents[second_doc].get("text", "")
                    
                    # Generate comparison
                    comparison = compare_documents(doc1_text, doc2_text, first_doc, second_doc)
                    if comparison:
                        st.markdown(comparison)
                        st.success("Comparison generated!")
                    else:
                        st.error("Failed to generate comparison.")
    
    # Chat input - outside of tabs
    if has_documents:
        if prompt := st.chat_input("Ask a question about your document(s)"):
            # Add user message to chat history
            st.session_state.chat_history.append({"role": "user", "content": prompt})
            
            # Get document context
            document_context = get_document_context(prompt)
            
            # Process user query
            response = process_user_query(prompt, document_context)
            
            # Add assistant response to chat history
            st.session_state.chat_history.append({"role": "assistant", "content": response})
            
            # Refresh the page to show the new messages
            st.rerun()

# Call main function at the end of the script
if __name__ == "__main__":
    main()